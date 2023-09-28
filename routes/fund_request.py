from config import app, db, login_manager, LAST_LEVEL, LEVEL_BEFORE_DGA_FOR_FDM, LEVEL_BEFORE_DGA_FOR_ODM, FIRST_LEVEL, DGA_LEVEL, LEVEL_NO_SIGNATURE_NEEDED, COMPLETED_STATUS, REJECTED_STATUS, fr_template_file_path, temp_folder, fund_doc_template_file_path
from models import FundRequest, EmployeeRequest, Employee, GradeZone, RequestApprovalDatetime, UserLevel, User
from .user import get_signature, get_initial
from werkzeug.utils import secure_filename
from flask import Blueprint, request, jsonify, request, send_file, session, send_from_directory
from flask_login import login_user, logout_user, login_required, current_user
from datetime import datetime
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches
from num2words import num2words
from docx2pdf import convert
import pythoncom
import uuid
import os
import locale


fund_request_bp = Blueprint("fund_request", __name__, url_prefix="/fund_requests")

locale.setlocale(locale.LC_ALL, 'fr_FR.UTF-8')

# Define the fund_request API routes
@fund_request_bp.route("/", methods=["POST"])
def create_fund_request():
    data = request.form

    # Extract data from request form and create a new FundRequest instance
    new_fund_request = FundRequest(
        itineraire=data.get('itineraire'),
        moyen_de_transport=data.get('moyen_de_transport'),
        zone_id=data.get('zone_id'),
        but_de_la_mission=data.get('but_de_la_mission'),
        date_debut=data.get('date_debut'),
        date_fin=data.get('date_fin'),
        domaine_activite_id=data.get('domaine_activite_id'),
        centre_de_cout_id=data.get('centre_de_cout_id'),
        nom_prenoms_chauffeur=data.get('nom_prenoms_chauffeur'),
        sup_hier=data.get('sup_hier'),
        dga=data.get('dga')
    )

    # Save the Fund Request Transaction
    db.session.add(new_fund_request)
    db.session.commit()
    new_fund_request.request_track_id = "TGNSCT"+str(new_fund_request.id).zfill(7)

    # Save the relations between the employees and the fund request

    employee_ids = data.get('participants').split(";")
    for employee_id in employee_ids:
        grade_id = Employee.query.get(employee_id).grade.id
        grade_zone = GradeZone.query.filter_by(grade_id = grade_id, zone_id = data.get('zone_id')).first()
        perdiem = grade_zone.montant_perdiem
        logement = grade_zone.montant_logement
        peage = grade_zone.montant_peage

        db.session.add(EmployeeRequest(
            fund_request_id=new_fund_request.id,
            employee_id=employee_id,
            montant_perdiem = perdiem * ((datetime.strptime(data.get('date_fin'), '%Y-%m-%d') - datetime.strptime(data.get('date_debut'), '%Y-%m-%d')).days+0.5),
            montant_logement = logement * ((datetime.strptime(data.get('date_fin'), '%Y-%m-%d') - datetime.strptime(data.get('date_debut'), '%Y-%m-%d')).days),
            montant_peage = peage
        ))

    db.session.commit()

    return jsonify({"fund_request_id": new_fund_request.id}), 201


@fund_request_bp.route("/<int:request_id>", methods=["GET"])
def get_fund_request(request_id):
    fund_request = FundRequest.query.get(request_id)

    if fund_request is None:
        return jsonify({"error": "Fund request not found"}), 404

    employee_ids = [employee_request.employee_id for employee_request in EmployeeRequest.query.filter_by(fund_request_id=request_id).all()]

    if employee_ids is None:
        return jsonify({"error": "Employees related to fund request not found"}), 404

    # Query the database to retrieve employees based on the list of IDs
    employees = Employee.query.filter(Employee.id.in_(employee_ids)).all()

    serialized_employees = []

    for employee in employees:
        emp_req = EmployeeRequest.query.filter_by(employee_id = employee.id, fund_request_id = request_id).first()
        emp_json = employee.serialize()
        emp_json['montant_perdiem'] = emp_req.montant_perdiem
        emp_json['montant_logement'] = emp_req.montant_logement
        emp_json['montant_peage'] = emp_req.montant_peage
        serialized_employees.append(emp_json)

    return jsonify(fund_request = fund_request.serialize(), employees = serialized_employees), 200

def get_fund_request_employees(request_id):
    fund_request = FundRequest.query.get(request_id)
    employee_ids = [employee_request.employee_id for employee_request in EmployeeRequest.query.filter_by(fund_request_id=request_id).all()]

    # Query the database to retrieve employees based on the list of IDs
    employees = Employee.query.filter(Employee.id.in_(employee_ids)).all()
    serialized_employees = []

    for employee in employees:
        emp_req = EmployeeRequest.query.filter_by(employee_id = employee.id, fund_request_id = request_id).first()
        emp_json = employee.serialize()
        emp_json['montant_perdiem'] = emp_req.montant_perdiem
        emp_json['montant_logement'] = emp_req.montant_logement
        emp_json['montant_peage'] = emp_req.montant_peage
        emp_json['montant_total'] = emp_req.montant_peage + emp_req.montant_logement + emp_req.montant_perdiem
        serialized_employees.append(emp_json)

    return fund_request, serialized_employees


@fund_request_bp.route("/", methods=["GET"])
@login_required
def get_all_fund_request():
    if int(current_user.user_level.level) == FIRST_LEVEL:
        fund_requests = FundRequest.query.filter_by(approval_level=current_user.user_level.level, superieur_hierarchique_username=current_user.username)
    elif int(current_user.user_level.level) == DGA_LEVEL:
        fund_requests = FundRequest.query.filter_by(approval_level=current_user.user_level.level, dga_username=current_user.username)
    else:
        fund_requests = FundRequest.query.filter_by(approval_level=current_user.user_level.level)
    
    if fund_requests.first() is None:
        return jsonify({"message": "No Funds request found."}), 200

    fund_requests_serialized = [fund_request.serialize() for fund_request in fund_requests]

    return jsonify(fund_requests_serialized), 200


@fund_request_bp.route("/approval_list/<int:request_id>", methods=["GET"])
def get_fund_request_list_approval(request_id):
    approvals = RequestApprovalDatetime.query.filter_by(fund_request_id=request_id)
    if approvals.first() is None:
        return jsonify({"message": "No Approval found."}), 204

    approvals_serialized = [approval.serialize() for approval in approvals]

    return jsonify(approvals_serialized), 200


@fund_request_bp.route("/approve", methods=["POST"])
@login_required
def approve_fund_request():
    data = request.form

    fund_request = FundRequest.query.get(data.get('fund_request_id'))

    if fund_request is None:
        return jsonify({"message": "Fund request not found"}), 404

    if (int(current_user.user_level.level) != int(fund_request.approval_level)):
        return jsonify({"message": "You're not authorized to approve this request."}), 404
        
    if (int(current_user.user_level.level) not in LEVEL_NO_SIGNATURE_NEEDED and current_user.signature_filename == None):
        return jsonify({"message": "You cannot approve. Set up your signature first"}), 404

    request_approval_dt = RequestApprovalDatetime(fund_request_id = data.get('fund_request_id'), level_id = current_user.user_level.id, user_id = current_user.id, approved_or_reject = 1)
    db.session.add(request_approval_dt)
    fund_request.approval_level+=1
    if int(current_user.user_level.level) == LAST_LEVEL:
        fund_request.status = COMPLETED_STATUS
    db.session.commit()

    return jsonify({"message": "Fund request approved."}), 200


@fund_request_bp.route("/reject", methods=["POST"])
@login_required
def reject_fund_request():
    data = request.form

    fund_request = FundRequest.query.get(data.get('fund_request_id'))

    if fund_request is None:
        return jsonify({"message": "Fund request not found"}), 404

    if int(current_user.user_level.level) != int(fund_request.approval_level):
        return jsonify({"message": "You're not authorized to reject this request."}), 404        

    request_approval_dt = RequestApprovalDatetime(fund_request_id = data.get('fund_request_id'), level_id = current_user.user_level.id, user_id = current_user.id, approved_or_reject = 0)
    db.session.add(request_approval_dt)
    fund_request.status = REJECTED_STATUS
    fund_request.approval_level = REJECTED_STATUS
    fund_request.reject_reason = data.get('reject_reason')
    db.session.commit()

    return jsonify({"message": "Fund request rejected."}), 200


@fund_request_bp.route("/print/<int:request_id>", methods=["GET"])
def print_fund_request(request_id):
    # Get parameters from Fund Request
    fund_request = FundRequest.query.get(request_id)
    duree_mission = (fund_request.date_fin - fund_request.date_debut).days+1

    data = {
        'ITINERAIRE': fund_request.itineraire,
        'BUTDELAMISSION': fund_request.but_de_la_mission,
        'DATEDEDEPART' : fund_request.date_debut.strftime("%d/%m/%Y"),
        'DATEDEFIN': fund_request.date_fin.strftime("%d/%m/%Y"),
        'DUREEMISSION': num2words(duree_mission, lang='fr').capitalize() + ' (' + str(duree_mission) +')',
        'MOYENDETRANSPORT': fund_request.moyen_de_transport,
        'NOMDUCHAUFFEUR': fund_request.nom_prenoms_chauffeur,
        'CENTREDECOUT': fund_request.cost_center.code,
        'ODM_NUMBER': str(fund_request.id).zfill(4),
        'ANNEE': datetime.now().year,
        'DATEDECREATION': fund_request.created_at.strftime("%d %B %Y"),
        'NOMAGENTS': '; '.join([Employee.query.get(employee_request.employee_id).nom_prenoms for employee_request in EmployeeRequest.query.filter_by(fund_request_id=request_id).all()]),
        'FONCTIONAGENTS':  '; '.join([Employee.query.get(employee_request.employee_id).fonction for employee_request in EmployeeRequest.query.filter_by(fund_request_id=request_id).all()]),
        'DGA': User.query.filter_by(username = fund_request.dga_username).first().firstname_lastname
    }

    doc = generate_fr_document(data, 'odm')
    filename = str(uuid.uuid4())
    filename_docx = filename + '.docx'
    filename_pdf = filename + '.pdf'
    temp_path = temp_folder + '/' + filename_docx
    temp_path_pdf = temp_folder + '/' + filename_pdf
    doc.save(temp_path)
    # Get approvers list and signatures or initial
    approvers = RequestApprovalDatetime.query.filter_by(fund_request_id=request_id, approved_or_reject='1').with_entities(RequestApprovalDatetime.user_id, RequestApprovalDatetime.level_id).all()

    initials = {}
    signature = None
    for i in approvers:
        current_level = UserLevel.query.filter_by(id=i[1]).first().level
        if int(current_level) == DGA_LEVEL:
            signature = get_signature(i[0])
        else:
            if i[0] != None and int(current_level) < DGA_LEVEL:
                initials[current_level] = get_initial(i[0])
    
    # Apply signatures and initial to saved Document
    document = Document(temp_path)
    image_paras = [i for i, p in enumerate(document.paragraphs) if "[haut]" in p.text]
    p = document.paragraphs[image_paras[0]]
    p.text = ""
    r = p.add_run()
    r.add_tab()
    r.add_tab()
    # add initial for levels before DGA signature
    for element in LEVEL_BEFORE_DGA_FOR_ODM:
        if str(element) in initials:
            r.add_picture(initials[str(element)], width=Inches(0.6), height=Inches(0.51))
            r.add_text("     ")
            os.remove(initials[str(element)])

    # Add signature below the ODM
    image_paras = [i for i, p in enumerate(document.paragraphs) if "[bas]" in p.text]
    p = document.paragraphs[image_paras[0]]
    p.text = ""
    r = p.add_run()
    for i in range(8):
        r.add_tab()
    
    # add signature for level 4
    if signature != None:
        r.add_picture(signature, width=Inches(0.6), height=Inches(0.51))
        os.remove(signature)
    
    document.save(temp_path)

    convert(temp_path, temp_path_pdf, pythoncom.CoInitialize())
    return send_from_directory(directory=os.getcwd()+temp_folder, path=filename_pdf, as_attachment=True), 200


@fund_request_bp.route("/generate_fund_doc/<int:request_id>", methods=["GET"])
@login_required
def generate_fund_doc(request_id):
    # Get parameters from Fund Request
    fund_request, employees = get_fund_request_employees(request_id)
    duree_mission = (fund_request.date_fin - fund_request.date_debut).days+1

    data = {
        'ITINERAIRE': fund_request.itineraire,
        'BUTDELAMISSION': fund_request.but_de_la_mission,
        'DATEDEDEPART' : fund_request.date_debut.strftime("%d/%m/%Y"),
        'DATEDEFIN': fund_request.date_fin.strftime("%d/%m/%Y"),
        'DUREEMISSION': num2words(duree_mission, lang='fr').capitalize() + ' (' + str(duree_mission) +')',
        'MOYENDETRANSPORT': fund_request.moyen_de_transport,
        'NOMDUCHAUFFEUR': fund_request.nom_prenoms_chauffeur,
        'CENTREDECOUT': fund_request.cost_center.code,
        'ODM_NUMBER': str(fund_request.id).zfill(4),
        'ANNEE': datetime.now().year,
        'DATEDECREATION': fund_request.created_at.strftime("%d %B %Y"),
        'NOMAGENTS': '; '.join([Employee.query.get(employee_request.employee_id).nom_prenoms for employee_request in EmployeeRequest.query.filter_by(fund_request_id=request_id).all()]),
        'FONCTIONAGENTS':  '; '.join([Employee.query.get(employee_request.employee_id).fonction for employee_request in EmployeeRequest.query.filter_by(fund_request_id=request_id).all()]),
        'DGA': User.query.filter_by(username = fund_request.dga_username).first().firstname_lastname,
        'employees': employees
    }

    doc = generate_fr_document(data, 'fdm')
    filename = str(uuid.uuid4())
    filename_docx = filename + '.docx'
    filename_pdf = filename + '.pdf'
    temp_path = temp_folder + '/' + filename_docx
    temp_path_pdf = temp_folder + '/' + filename_pdf
    doc.save(temp_path)
    # Get approvers list and signatures or initial
    approvers = RequestApprovalDatetime.query.filter_by(fund_request_id=request_id, approved_or_reject='1').with_entities(RequestApprovalDatetime.user_id, RequestApprovalDatetime.level_id).all()

    initials = {}
    signature = None
    for i in approvers:
        current_level = UserLevel.query.filter_by(id=i[1]).first().level
        if int(current_level) == DGA_LEVEL:
            signature = get_signature(i[0])
        else:
            if i[0] != None and int(current_level) > DGA_LEVEL:
                initials[current_level] = get_initial(i[0])
    
    # Apply signatures and initial to saved Document
    document = Document(temp_path)
    image_paras = [i for i, p in enumerate(document.paragraphs) if "[haut]" in p.text]
    p = document.paragraphs[image_paras[0]]
    p.text = ""
    r = p.add_run()
    r.add_tab()
    r.add_tab()
    # add initial for levels before DGA signature
    for element in LEVEL_BEFORE_DGA_FOR_FDM:
        if str(element) in initials:
            r.add_picture(initials[str(element)], width=Inches(0.6), height=Inches(0.51))
            r.add_text("     ")
            os.remove(initials[str(element)])

    # Add signature below the ODM
    image_paras = [i for i, p in enumerate(document.paragraphs) if "[bas]" in p.text]
    p = document.paragraphs[image_paras[0]]
    p.text = ""
    r = p.add_run()
    for i in range(8):
        r.add_tab()
    
    # add signature for level 4
    if signature != None:
        r.add_picture(signature, width=Inches(0.6), height=Inches(0.51))
        os.remove(signature)
    
    document.save(temp_path)
    
    convert(temp_path, temp_path_pdf, pythoncom.CoInitialize())
    return send_from_directory(directory=os.getcwd()+temp_folder, path=filename_pdf, as_attachment=True), 200


# ODM is ordre de mission and FDM is frais de mission
def generate_fr_document(data, type):
    if type=='odm':
        doc = DocxTemplate(fr_template_file_path)
    elif type=='fdm':
        doc = DocxTemplate(fund_doc_template_file_path)
    doc.render(data)
    return doc